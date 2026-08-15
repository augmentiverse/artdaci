from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\Users\nimdo\Documents\GitHub\artdaci")
SOURCE = ROOT / "artifacts" / "ARTDACI_Prototype_Papier_35_Pages_FR.docx"
OUT = ROOT / "artifacts" / "ARTDACI_Prototype_Papier_Enrichi_39_Pages_FR.docx"

NAVY = RGBColor(18, 43, 64)
GOLD = RGBColor(186, 137, 62)
MUTED = RGBColor(94, 105, 116)

pages = [
    {
        "artist": "Léonard de Vinci",
        "title": "Un savoir construit par l’observation",
        "intro": "Chez Léonard, peindre, dessiner et comprendre relèvent d’un même mouvement. Son œuvre, peu abondante mais extraordinairement influente, naît d’une curiosité qui traverse le corps humain, la lumière, l’eau, les machines et le paysage.",
        "sections": [
            ("Florence : apprendre dans un atelier collectif", "Né en 1452 à Vinci, Léonard est formé à Florence dans l’atelier d’Andrea del Verrocchio. Il y rencontre peinture, sculpture, dessin, perspective et techniques de fabrication. Les œuvres de jeunesse montrent déjà son attention aux transitions lumineuses, aux gestes et aux expressions. L’Annonciation et le Baptême du Christ témoignent d’un apprentissage où l’observation de la nature compte autant que la maîtrise du métier."),
            ("Milan : l’artiste au service d’une cour", "À partir des années 1480, il travaille pour Ludovic Sforza. Il conçoit décors, fêtes, projets d’ingénierie et monuments, tout en réalisant La Cène. Cette peinture murale transforme le dernier repas du Christ en drame psychologique : chaque apôtre réagit à l’annonce de la trahison. Le choix d’une technique expérimentale, différente de la fresque traditionnelle, permet des effets subtils mais fragilise rapidement l’œuvre."),
            ("Dessiner pour penser", "Ses carnets associent mots, schémas et études. Anatomie, mouvement des fluides, botanique, optique ou mécanique y sont examinés par comparaison. L’Homme de Vitruve ne se réduit pas à une icône : il traduit une recherche sur les proportions et la relation entre corps, géométrie et architecture. Plusieurs machines dessinées n’ont jamais été construites ; leur importance tient surtout à la méthode de visualisation et de résolution des problèmes."),
            ("Une œuvre ouverte", "La Joconde, commencée à Florence et retravaillée durant plusieurs années, résume cette pensée des relations : visage et paysage, matière et atmosphère, présence et incertitude. Les documents ne permettent pas de trancher toutes les questions sur sa genèse. Cette part d’inachevé n’est pas un défaut de connaissance : elle invite à distinguer ce qui est attesté, probable ou interprété."),
        ],
        "note": "À retenir — Chez Léonard, le dessin est à la fois mémoire, expérience et outil d’invention.",
    },
    {
        "artist": "Vincent van Gogh",
        "title": "La couleur comme expérience vécue",
        "intro": "La carrière artistique de Van Gogh tient en une décennie environ. Sa peinture évolue très vite, nourrie par le dessin, la littérature, les échanges avec son frère Theo et l’étude directe des œuvres anciennes comme modernes.",
        "sections": [
            ("Des Pays-Bas à Paris", "Après plusieurs emplois et une période de prédication dans le Borinage, Van Gogh décide de devenir artiste. Ses premières œuvres privilégient les travailleurs, les tisserands et les paysans, dans une gamme sombre. Les Mangeurs de pommes de terre cherche une vérité sociale et expressive plutôt qu’une beauté idéale. À Paris, dès 1886, la rencontre de l’impressionnisme, du néo-impressionnisme et des estampes japonaises éclaircit sa palette et libère sa touche."),
            ("Arles : inventer un atelier du Sud", "Installé à Arles en 1888, il espère créer une communauté d’artistes. La lumière méridionale intensifie ses jaunes, bleus et verts. Les Tournesols, La Chambre et les paysages de vergers montrent qu’il simplifie les formes et utilise la couleur pour construire l’espace autant que pour exprimer une sensation. La cohabitation avec Paul Gauguin se termine dans une crise grave ; l’épisode de l’oreille demeure entouré de détails difficiles à établir avec certitude."),
            ("Saint-Rémy et Auvers", "Lors de son séjour volontaire à l’asile de Saint-Rémy, Van Gogh continue de travailler entre crises, d’après le paysage et d’après des reproductions. La Nuit étoilée transforme l’observation en vision rythmique : ciel, cyprès et village sont unifiés par le mouvement de la touche. À Auvers-sur-Oise, il peint avec une grande intensité sous le regard du docteur Gachet. Il meurt en juillet 1890, deux jours après une blessure par balle généralement considérée comme volontaire."),
            ("Une postérité construite", "Sa reconnaissance n’est pas un miracle instantané. Theo soutient matériellement et intellectuellement son travail ; après leurs morts, Johanna van Gogh-Bonger classe les lettres, organise des expositions et favorise la diffusion des œuvres. Les lettres permettent de suivre ses choix de couleurs et ses ambitions, mais elles ne doivent pas réduire toute la peinture à la seule biographie."),
        ],
        "note": "À retenir — L’intensité de Van Gogh repose sur une pratique réfléchie : dessin, contrastes colorés, rythme et séries.",
    },
    {
        "artist": "Johannes Vermeer",
        "title": "L’art de ralentir le regard",
        "intro": "Vermeer laisse un corpus très restreint, généralement évalué à une trentaine de peintures reconnues. Ses scènes semblent simples, mais elles reposent sur une organisation rigoureuse de la lumière, des plans, des objets et des gestes.",
        "sections": [
            ("Delft : famille, métier et marché", "Né en 1632, Vermeer grandit dans une famille liée au commerce de l’art et aux métiers du textile. Il épouse Catharina Bolnes et entre dans la guilde de Saint-Luc en 1653. Il en devient plusieurs fois responsable, signe d’une réelle considération locale. Son activité se déroule presque entièrement à Delft. Une production lente et un cercle d’acheteurs limité expliquent en partie le faible nombre d’œuvres conservées."),
            ("Construire la lumière", "Dans ses intérieurs, une fenêtre placée à gauche distribue souvent la lumière. Les sols carrelés, tables, chaises, cartes et tentures organisent la profondeur. Vermeer combine contours nets et zones adoucies ; de petits accents lumineux suggèrent reflets et textures. Il emploie des pigments coûteux, notamment l’outremer naturel. L’hypothèse d’un recours à la chambre noire est plausible pour comprendre certains effets optiques, mais aucun document ne prouve qu’il ait copié mécaniquement une projection."),
            ("Des récits suspendus", "Une femme lit, verse du lait, tient une balance ou regarde le spectateur. L’action est minimale, pourtant chaque objet peut orienter l’interprétation. Les tableaux, cartes, instruments et lettres introduisent des thèmes de connaissance, d’amour, de vanité ou de mesure. Il faut toutefois éviter un dictionnaire automatique des symboles : leur sens dépend de la composition entière et du contexte culturel."),
            ("Oubli, redécouverte et expertise", "Après sa mort en 1675, Vermeer est longtemps confondu avec d’autres peintres. Sa redécouverte au XIXe siècle contribue à former sa réputation moderne. Au XXe siècle, les faux de Han van Meegeren rappellent que l’attribution repose sur la provenance, l’étude stylistique et l’analyse scientifique des matériaux. Les recherches techniques actuelles révèlent aussi des transformations sous la surface et restituent le processus de création."),
        ],
        "note": "À retenir — Vermeer rend visible l’attention elle-même : regarder devient le véritable sujet de la scène.",
    },
    {
        "artist": "Claude Monet",
        "title": "Peindre les variations du visible",
        "intro": "Monet ne cherche pas seulement à représenter un lieu. Il observe comment l’heure, le temps, l’air et les reflets transforment ce lieu. Sa peinture fait de la perception un sujet à part entière.",
        "sections": [
            ("Le Havre et l’apprentissage du plein air", "Né à Paris en 1840 et élevé au Havre, Monet commence par la caricature. Eugène Boudin l’encourage à peindre dehors et à étudier directement le ciel et la mer. Cette pratique lui apprend à travailler vite, à comparer les effets atmosphériques et à accorder au paysage moderne — ports, gares, ponts, loisirs — la dignité d’un sujet majeur."),
            ("1874 : un titre devient un mouvement", "Impression, soleil levant représente le port du Havre dans la brume. Exposé en 1874 avec des œuvres d’artistes indépendants, le tableau inspire au critique Louis Leroy le mot « impressionnistes », d’abord moqueur. Le terme est ensuite adopté. La peinture n’est pourtant pas une esquisse négligée : les silhouettes, les reflets et le disque solaire sont soigneusement distribués pour produire profondeur et vibration."),
            ("Couleur, lumière et distance", "Monet juxtapose des touches visibles et limite les contours fermés. Les ombres reçoivent des couleurs plutôt qu’un simple noir. Dans Impression, soleil levant, le contraste entre l’orange du soleil et les bleus gris du port attire fortement l’œil, même si leurs valeurs lumineuses sont proches. De près, la surface paraît fragmentée ; à distance, la scène s’unifie. Cette double lecture engage activement le spectateur."),
            ("Séries et jardin de Giverny", "À partir des années 1890, Monet peint un même motif sous des conditions changeantes : Meules, Peupliers, Cathédrales de Rouen. À Giverny, le bassin aux nymphéas devient un laboratoire de reflets, de profondeur et de cadrage. Les grandes décorations de l’Orangerie enveloppent le regard et annoncent certaines ambitions de l’abstraction, sans cesser d’être fondées sur l’expérience du jardin."),
        ],
        "note": "À retenir — Une série de Monet ne répète pas un motif : elle montre que le visible ne cesse de changer.",
    },
]

def paragraph_before(anchor, text="", style=None):
    p_el = OxmlElement("w:p")
    anchor._p.addprevious(p_el)
    from docx.text.paragraph import Paragraph
    p = Paragraph(p_el, anchor._parent)
    if style:
        p.style = style
    if text:
        p.add_run(text)
    return p

def shade(p, fill="EEF3F5"):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)

doc = Document(SOURCE)
anchors = [p for p in doc.paragraphs if p.text == "CHEF-D’ŒUVRE"]
if len(anchors) != 4:
    raise RuntimeError(f"Quatre ancres attendues, {len(anchors)} trouvées")

for anchor, data in zip(anchors, pages):
    # Reuse the explicit page break already placed before the masterpiece page.
    break_p = anchor._p.getprevious()
    from docx.text.paragraph import Paragraph
    break_para = Paragraph(break_p, anchor._parent)

    pb = paragraph_before(break_para)
    pb.add_run().add_break(WD_BREAK.PAGE)

    kicker = paragraph_before(break_para, f"APPROFONDISSEMENT · {data['artist'].upper()}")
    kicker.paragraph_format.space_after = Pt(6)
    r = kicker.runs[0]
    r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = GOLD

    title = paragraph_before(break_para, data["title"], "Heading 1")
    title.paragraph_format.space_after = Pt(7)

    intro = paragraph_before(break_para, data["intro"])
    intro.paragraph_format.space_after = Pt(7)
    intro.paragraph_format.line_spacing = 1.08
    for r in intro.runs:
        r.italic = True; r.font.color.rgb = MUTED

    for heading, body in data["sections"]:
        h = paragraph_before(break_para, heading, "Heading 2")
        h.paragraph_format.space_before = Pt(5)
        h.paragraph_format.space_after = Pt(2)
        p = paragraph_before(break_para, body)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.06
        for r in p.runs:
            r.font.size = Pt(9.2)

    note = paragraph_before(break_para, data["note"])
    note.paragraph_format.space_before = Pt(6)
    note.paragraph_format.space_after = Pt(0)
    note.paragraph_format.left_indent = Pt(8)
    note.paragraph_format.right_indent = Pt(8)
    shade(note)
    for r in note.runs:
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = NAVY

# Mention the enrichment in the contents without introducing fragile page numbers.
for p in doc.paragraphs:
    if p.text.startswith("Chaque chapitre suit le même rythme"):
        p.text = ("Chaque chapitre suit le même rythme : portrait et biographie, parcours et méthode, "
                  "approfondissement, chef-d’œuvre, lecture guidée, analyse détaillée et prolongements numériques facultatifs.")
        break

for p in doc.paragraphs:
    if p.text.startswith("Rapport documentaire fourni"):
        p.text = ("Documents de travail fournis : 4painters-report.docx (10 août 2026) ; "
                  "page pour chaque peintre.docx (contenu traduit, vérifié, condensé et adapté à la maquette).")
        break

doc.core_properties.title = "ARTDACI — Prototype papier enrichi, 39 pages"
doc.core_properties.subject = "Livre imprimé autonome consacré à Léonard de Vinci, Vincent van Gogh, Johannes Vermeer et Claude Monet"
doc.core_properties.comments = "Édition enrichie à partir du document « page pour chaque peintre.docx » ; contenu traduit, condensé et harmonisé."
doc.save(OUT)
print(OUT)
