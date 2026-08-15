from pathlib import Path
import sys
from docx import Document

source = Path(r"C:\Users\nimdo\Downloads\ARTDACI Master Production Plan.docx")
sys.stdout.reconfigure(encoding="utf-8")
doc = Document(source)
print(f"sections={len(doc.sections)} paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} inline_shapes={len(doc.inline_shapes)}")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        print(f"P{i:03d}\t[{p.style.name}]\t{text}")
for ti, table in enumerate(doc.tables):
    print(f"TABLE {ti} rows={len(table.rows)} cols={len(table.columns)} style={table.style.name if table.style else ''}")
    for ri, row in enumerate(table.rows):
        print(f"R{ri:02d}\t" + " | ".join(cell.text.replace("\n", " / ") for cell in row.cells))
for si, section in enumerate(doc.sections):
    print(f"SECTION {si}: page={section.page_width}x{section.page_height} margins={section.top_margin},{section.right_margin},{section.bottom_margin},{section.left_margin}")
    for part_name, part in (("HEADER", section.header), ("FOOTER", section.footer)):
        for p in part.paragraphs:
            if p.text.strip(): print(f"{part_name}\t{p.text}")
