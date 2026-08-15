from pathlib import Path
from copy import deepcopy
from PIL import Image, ImageOps, ImageEnhance
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\Users\nimdo\Documents\GitHub\artdaci")
SOURCE = ROOT / "artifacts" / "ARTDACI_Livre_imprime_Quatre_Maitres_FR.docx"
OUT = ROOT / "artifacts" / "ARTDACI_Prototype_Papier_35_Pages_FR.docx"
MEDIA = ROOT / "tmp" / "prototype-media"
MEDIA.mkdir(parents=True, exist_ok=True)

NAVY = RGBColor(18, 43, 64)
GOLD = RGBColor(186, 137, 62)
MUTED = RGBColor(94, 105, 116)

detail_sources = {
    "Éloignez-vous puis rapprochez-vous : le sourire paraît se modifier.": (
        ROOT / "assets/paintings/Da Vinci/mona-lisa/images/mona-lisa.jpg",
        "Détail d’observation : visage, mains et paysage de La Joconde.",
    ),
    "Observez comment le vêtement sombre stabilise la composition.": (
        ROOT / "assets/paintings/van-gogh/tableaux/Autoportrait_VanGogh.png",
        "Détail d’observation : direction de la touche et contrastes de l’Autoportrait.",
    ),
    "Imaginez le rideau vert désormais presque noir.": (
        ROOT / "assets/paintings/Vermeer/Tableaux/Girl with a Pearl Earring_Vermeer.png",
        "Détail d’observation : regard, lèvres, turban et reflet de la perle.",
    ),
    "Observez comment deux barques suffisent à donner l’échelle et la profondeur.": (
        ROOT / "assets/paintings/monet/Tableaux/Impression-Sunrise_Monet.png",
        "Détail d’observation : soleil, reflets, barques et brume du port du Havre.",
    ),
}

portfolios = {
    "Léonard de Vinci — quatre œuvres en regard": [
        ROOT / "assets/paintings/Da Vinci/Tableaux/Mana Lisa_DaVici.png",
        ROOT / "assets/paintings/Da Vinci/Tableaux/The Lady with an Ermine_DaVinci.png",
        ROOT / "assets/paintings/Da Vinci/Tableaux/The Annunciation_DaVinci.png",
        ROOT / "assets/paintings/Da Vinci/Tableaux/The Last Supper_DaVinci.png",
    ],
    "Vincent van Gogh — quatre œuvres en regard": [
        ROOT / "assets/paintings/van-gogh/tableaux/Autoportrait_VanGogh.png",
        ROOT / "assets/paintings/van-gogh/tableaux/The Starry Night_VanGogh.png",
        ROOT / "assets/paintings/van-gogh/tableaux/Tournesols_VanGogh.png",
        ROOT / "assets/paintings/van-gogh/tableaux/The Bedroom_VanGogh.png",
    ],
    "Johannes Vermeer — quatre œuvres en regard": [
        ROOT / "assets/paintings/Vermeer/Tableaux/Girl with a Pearl Earring_Vermeer.png",
        ROOT / "assets/paintings/Vermeer/Tableaux/The Milkmaid_Vermeer.png",
        ROOT / "assets/paintings/Vermeer/Tableaux/View of Delft_Vermeer.png",
        ROOT / "assets/paintings/Vermeer/Tableaux/The Astronomer_Vermeer.png",
    ],
    "Claude Monet — quatre œuvres en regard": [
        ROOT / "assets/paintings/monet/Tableaux/Impression-Sunrise_Monet.png",
        ROOT / "assets/paintings/monet/Tableaux/Woman with a parasol_Monet.png",
        ROOT / "assets/paintings/monet/Tableaux/Water Lilies_Monet.png",
        ROOT / "assets/paintings/monet/Tableaux/The Japanese Bridge_Monet.png",
    ],
}

portfolio_captions = {
    "Léonard de Vinci — quatre œuvres en regard": "De gauche à droite, en haut puis en bas : La Joconde ; La Dame à l’hermine ; L’Annonciation ; La Cène.",
    "Vincent van Gogh — quatre œuvres en regard": "Autoportrait ; La Nuit étoilée ; Tournesols ; La Chambre à Arles.",
    "Johannes Vermeer — quatre œuvres en regard": "La Jeune Fille à la perle ; La Laitière ; Vue de Delft ; L’Astronome.",
    "Claude Monet — quatre œuvres en regard": "Impression, soleil levant ; Femme à l’ombrelle ; Nymphéas ; Le Pont japonais.",
}

def crop_for_page(src: Path, out: Path, size=(1500, 920)):
    with Image.open(src) as im:
        im = im.convert("RGB")
        fitted = ImageOps.fit(im, size, method=Image.Resampling.LANCZOS, centering=(0.5, 0.45))
        fitted = ImageEnhance.Contrast(fitted).enhance(1.03)
        fitted.save(out, quality=93)

def montage(paths, out: Path):
    cell = (720, 500)
    gap = 24
    canvas = Image.new("RGB", (cell[0]*2 + gap*3, cell[1]*2 + gap*3), (18, 43, 64))
    for idx, src in enumerate(paths):
        with Image.open(src) as im:
            tile = ImageOps.fit(im.convert("RGB"), cell, method=Image.Resampling.LANCZOS)
            x = gap + (idx % 2) * (cell[0] + gap)
            y = gap + (idx // 2) * (cell[1] + gap)
            canvas.paste(tile, (x, y))
    canvas.save(out, quality=92)

def set_alt(inline_shape, text):
    inline_shape._inline.docPr.set("descr", text)
    inline_shape._inline.docPr.set("title", text[:80])

def insert_after(paragraph, text=None, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_p, paragraph._parent)
    if style:
        p.style = style
    if text:
        p.add_run(text)
    return p

def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text.upper())
    r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = GOLD
    return p

def add_title(doc, text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_after = Pt(10)
    return p

def add_body(doc, text, after=7):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.16
    return p

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)

doc = Document(SOURCE)

# Add the new prototype supplement to the contents page.
for p in doc.paragraphs:
    if p.text.startswith("Sources, crédits et repères"):
        newp = insert_after(p, "Cahier transversal — chronologie, portfolios, techniques et test de lecture", "List Bullet")
        if p._p.pPr is not None:
            if newp._p.pPr is not None:
                newp._p.remove(newp._p.pPr)
            newp._p.insert(0, deepcopy(p._p.pPr))
        newp.paragraph_format.space_after = Pt(3)
        break

# Monet's observation instructions begin a deliberate detail page. Remove the
# old break after the bullets so the heading, prompts, image and caption stay together.
observation_heads = [p for p in doc.paragraphs if p.text.strip() == "À regarder devant l’image"]
monet_head = observation_heads[-1]
monet_head.paragraph_format.page_break_before = True
monet_last_prompt = next(p for p in doc.paragraphs if p.text.strip() == "Observez comment deux barques suffisent à donner l’échelle et la profondeur.")
old_break = monet_last_prompt._p.getnext()
if old_break is not None and 'w:type="page"' in old_break.xml:
    old_break.getparent().remove(old_break)
analysis_kickers = [p for p in doc.paragraphs if p.text.strip() == "ANALYSE APPROFONDIE"]
analysis_kickers[-1].paragraph_format.page_break_before = True

# Turn the four observation pages into intentional full visual spreads.
for anchor, (src, caption) in detail_sources.items():
    target = next(p for p in doc.paragraphs if p.text.strip() == anchor)
    img_path = MEDIA / (src.stem.replace(" ", "-") + "-detail.jpg")
    crop_for_page(src, img_path)
    pic_p = insert_after(target)
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.space_before = Pt(12)
    pic_p.paragraph_format.space_after = Pt(5)
    shape = pic_p.add_run().add_picture(str(img_path), width=Inches(5.8))
    set_alt(shape, caption)
    cap_p = insert_after(pic_p, caption)
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_after = Pt(0)
    for r in cap_p.runs:
        r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = MUTED

# Eight-page print-only supplement: the prototype reaches 34 pages.
doc.add_page_break()
add_kicker(doc, "Cahier transversal")
add_title(doc, "Quatre manières de regarder")
portrait_montage = MEDIA / "quatre-maitres.jpg"
montage([
    ROOT / "assets/peintres/photos des peintres/da-vinci-f.png",
    ROOT / "assets/peintres/photos des peintres/van-goh-f.png",
    ROOT / "assets/peintres/photos des peintres/vermeer-f.png",
    ROOT / "assets/peintres/photos des peintres/monet-f.png",
], portrait_montage)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
shape = p.add_run().add_picture(str(portrait_montage), width=Inches(5.75)); set_alt(shape, "Portraits de Léonard de Vinci, Vincent van Gogh, Johannes Vermeer et Claude Monet.")
add_body(doc, "Ces quatre artistes ne forment pas une école unique. Ils permettent de comparer quatre ambitions : comprendre le monde par le dessin, traduire l’intensité intérieure, organiser la lumière dans l’espace domestique et saisir les variations fugitives de l’atmosphère.")
add_bullets(doc, [
    "Léonard : observer, construire et relier art et connaissance.",
    "Van Gogh : rendre visible l’énergie du geste et de la couleur.",
    "Vermeer : mettre en scène le silence, la lumière et l’attention.",
    "Monet : peindre les conditions changeantes de la perception.",
])

doc.add_page_break()
add_kicker(doc, "Repères")
add_title(doc, "Une chronologie en quatre temps")
chronology = [
    ("1452–1519", "Léonard de Vinci", "Renaissance italienne", "Perspective, anatomie, sfumato, pensée par le dessin."),
    ("1632–1675", "Johannes Vermeer", "Siècle d’or néerlandais", "Intérieurs, lumière latérale, pigments précieux, scènes suspendues."),
    ("1840–1926", "Claude Monet", "Impressionnisme", "Plein air, séries, couleur optique, changements de lumière."),
    ("1853–1890", "Vincent van Gogh", "Postimpressionnisme", "Touche expressive, contrastes complémentaires, peinture comme expérience vécue."),
]
for dates, artist, period, insight in chronology:
    h = doc.add_heading(f"{dates} · {artist}", level=2)
    h.paragraph_format.space_before = Pt(8); h.paragraph_format.space_after = Pt(3)
    p = add_body(doc, f"{period}. {insight}", after=5)
add_body(doc, "À retenir : les artistes ne se succèdent pas selon une progression simple. Chacun répond aux savoirs, aux marchés, aux matériaux et aux attentes de son époque. La comparaison sert à distinguer des problèmes de peinture, non à établir un classement.")

portfolio_notes = {
    "Léonard de Vinci — quatre œuvres en regard": ["Cherchez les transitions plutôt que les contours.", "Comparez le rôle des mains et des regards.", "Repérez la géométrie qui stabilise les figures."],
    "Vincent van Gogh — quatre œuvres en regard": ["Suivez le rythme de la touche.", "Observez les contrastes entre couleurs complémentaires.", "Comparez espace intérieur, paysage et autoportrait."],
    "Johannes Vermeer — quatre œuvres en regard": ["Identifiez la source de lumière.", "Regardez les bords nets puis les contours adoucis.", "Mesurez la place du vide et du silence."],
    "Claude Monet — quatre œuvres en regard": ["Comparez les touches de près et l’unité de loin.", "Repérez les reflets et les couleurs de l’ombre.", "Imaginez le changement de lumière quelques minutes plus tard."],
}
for idx, (title, paths) in enumerate(portfolios.items(), start=1):
    doc.add_page_break()
    add_kicker(doc, f"Portfolio {idx} sur 4")
    add_title(doc, title)
    out = MEDIA / f"portfolio-{idx}.jpg"
    montage(paths, out)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(5)
    shape = p.add_run().add_picture(str(out), width=Inches(5.9)); set_alt(shape, portfolio_captions[title])
    cap = doc.add_paragraph(portfolio_captions[title]); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; cap.paragraph_format.space_after = Pt(8)
    for r in cap.runs: r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = MUTED
    add_body(doc, "Exercice de regard", after=3).runs[0].bold = True
    add_bullets(doc, portfolio_notes[title])

doc.add_page_break()
add_kicker(doc, "Boîte à outils")
add_title(doc, "Techniques et vocabulaire essentiels")
glossary = [
    ("Composition", "Organisation des formes, des masses, des directions et des vides dans l’image."),
    ("Glacis", "Couche picturale très mince et translucide qui modifie profondeur, teinte ou luminosité."),
    ("Sfumato", "Transitions fondues entre ombre et lumière, sans contour brutal ; procédé associé à Léonard."),
    ("Empâtement", "Peinture posée en épaisseur, dont le relief conserve la trace de l’outil ou du pinceau."),
    ("Couleurs complémentaires", "Couleurs opposées qui renforcent mutuellement leur intensité, comme orange et bleu."),
    ("Valeur tonale", "Degré de clarté ou d’obscurité d’une couleur, indépendamment de sa teinte."),
    ("Perspective atmosphérique", "Effet de profondeur obtenu par l’adoucissement des contrastes et le bleuissement des lointains."),
    ("Provenance", "Historique documenté des propriétaires, ventes, collections et déplacements d’une œuvre."),
]
for term, definition in glossary:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(term + " — "); r.bold = True; r.font.color.rgb = NAVY
    p.add_run(definition)
add_body(doc, "Méthode en trois gestes", after=4).runs[0].bold = True
add_bullets(doc, ["Décrire d’abord ce qui est visible.", "Relier ensuite l’effet aux choix de matière et de composition.", "Distinguer enfin le fait établi, l’hypothèse et l’interprétation."])

doc.add_page_break()
add_kicker(doc, "Test du prototype")
add_title(doc, "Carnet du lecteur")
add_body(doc, "Cette page permet d’évaluer le livre sans application. Répondez après une lecture libre, sans scanner de QR code et sans ouvrir la galerie virtuelle.")
questions = [
    "Quelle œuvre vous a retenu le plus longtemps, et pourquoi ?",
    "Citez une technique propre à chacun des quatre artistes.",
    "Avez-vous compris le contexte et l’importance des quatre œuvres principales sans contenu numérique ?",
    "Quelle page vous a semblé la plus claire ? La plus dense ?",
    "Une information essentielle vous a-t-elle manqué ?",
    "Les reproductions sont-elles assez grandes pour observer les détails indiqués ?",
    "Le rythme entre images, textes et encadrés vous paraît-il confortable ?",
]
for i, question in enumerate(questions, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{i}. {question}"); r.bold = True; r.font.color.rgb = NAVY
    line = doc.add_paragraph("________________________________________________________________________________")
    line.paragraph_format.space_after = Pt(4)
    for r in line.runs: r.font.color.rgb = RGBColor(185, 188, 191); r.font.size = Pt(8)
add_body(doc, "Principe de validation : si le lecteur peut expliquer les quatre artistes et leurs œuvres majeures en s’appuyant uniquement sur ces pages, le prototype remplit sa fonction de livre autonome.")

# Keep headings and figures together where Word supports it.
for p in doc.paragraphs:
    if p.style.name.startswith("Heading"):
        p.paragraph_format.keep_with_next = True
    p.paragraph_format.widow_control = True

doc.core_properties.title = "ARTDACI — Prototype papier de 35 pages"
doc.core_properties.subject = "Livre autonome consacré à Léonard de Vinci, Van Gogh, Vermeer et Monet"
doc.core_properties.keywords = "ARTDACI, livre imprimé, art, Léonard de Vinci, Van Gogh, Vermeer, Monet"
doc.save(OUT)
print(OUT)
