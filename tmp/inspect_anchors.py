from docx import Document

d = Document(r"artifacts\ARTDACI_Prototype_Papier_35_Pages_FR.docx")
inds = [i for i, p in enumerate(d.paragraphs) if p.text == "CHEF-D’ŒUVRE"]
print(inds)
for i in inds:
    print("\nANCHOR", i)
    for j in range(i - 5, i + 3):
        p = d.paragraphs[j]
        print(j, repr(p.text), "breakbefore", p.paragraph_format.page_break_before,
              "xmlbreak", 'w:type="page"' in p._p.xml)
