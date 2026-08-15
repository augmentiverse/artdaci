from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(r"C:\Users\nimdo\Documents\GitHub\artdaci\artifacts\ARTDACI_Master_Production_Plan_Ameliore_FR.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = "10283F"
BLUE = "245E78"
GOLD = "B28A3B"
INK = "202B33"
MUTED = "5D6B75"
PALE = "EAF0F4"
PALE_GOLD = "F5EFE2"
WHITE = "FFFFFF"
GREEN = "2F6D55"
RED = "8B3A3A"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.72)
sec.bottom_margin = Inches(0.72)
sec.left_margin = Inches(0.78)
sec.right_margin = Inches(0.78)
sec.header_distance = Inches(0.32)
sec.footer_distance = Inches(0.35)

def set_font(run, name="Aptos", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color is not None: run.font.color.rgb = RGBColor.from_string(color)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(10.2)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.16
for name, size, color, before, after in [
    ("Heading 1", 16, NAVY, 14, 7),
    ("Heading 2", 12.5, BLUE, 10, 5),
    ("Heading 3", 10.8, NAVY, 7, 3),
]:
    s = styles[name]
    s.font.name = "Aptos Display"
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True
for name in ("List Bullet", "List Number"):
    s = styles[name]
    s.font.name = "Aptos"
    s.font.size = Pt(10.2)
    s.paragraph_format.left_indent = Inches(0.32)
    s.paragraph_format.first_line_indent = Inches(-0.18)
    s.paragraph_format.space_after = Pt(3)
    s.paragraph_format.line_spacing = 1.12

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")

def set_cell_text(cell, text, bold=False, color=INK, size=9.2, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.03
    r = p.add_run(str(text))
    set_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell_margins(cell)

def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

def prevent_row_split(row):
    trPr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    trPr.append(cant_split)

def table(rows, widths, header=True, font_size=9.2):
    t = doc.add_table(rows=1, cols=len(rows[0]))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.style = "Table Grid"
    for i, w in enumerate(widths):
        t.columns[i].width = Inches(w)
    for i, val in enumerate(rows[0]):
        set_cell_text(t.rows[0].cells[i], val, bold=True, color=WHITE, size=font_size)
        shade(t.rows[0].cells[i], NAVY)
    set_repeat_header(t.rows[0])
    prevent_row_split(t.rows[0])
    for ri, row in enumerate(rows[1:], 1):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=font_size)
            if ri % 2 == 0: shade(cells[i], "F5F7F9")
        prevent_row_split(t.rows[-1])
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return t

def add_p(text="", bold=False, italic=False, color=INK, align=None, size=None, after=None):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    if after is not None: p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_font(r, size=size or 10.2, bold=bold, italic=italic, color=color)
    return p

def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

def add_numbers(items):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType"); multi.set(qn("w:val"), "singleLevel"); abstract.append(multi)
    lvl = OxmlElement("w:lvl"); lvl.set(qn("w:ilvl"), "0"); abstract.append(lvl)
    start = OxmlElement("w:start"); start.set(qn("w:val"), "1"); lvl.append(start)
    fmt = OxmlElement("w:numFmt"); fmt.set(qn("w:val"), "decimal"); lvl.append(fmt)
    text = OxmlElement("w:lvlText"); text.set(qn("w:val"), "%1."); lvl.append(text)
    suff = OxmlElement("w:suff"); suff.set(qn("w:val"), "space"); lvl.append(suff)
    ppr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "540"); ind.set(qn("w:hanging"), "280"); ppr.append(ind)
    lvl.append(ppr)
    numbering.append(abstract)
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId"); abstract_ref.set(qn("w:val"), str(abstract_id)); num.append(abstract_ref)
    numbering.append(num)
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12
        ppr = p._p.get_or_add_pPr()
        numpr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
        numid = OxmlElement("w:numId"); numid.set(qn("w:val"), str(num_id))
        numpr.extend([ilvl, numid]); ppr.append(numpr)
        p.add_run(item)

def heading(text, level=1):
    return doc.add_heading(text, level=level)

def callout(title, body, fill=PALE_GOLD):
    t = doc.add_table(rows=1, cols=1)
    prevent_row_split(t.rows[0])
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Inches(6.75)
    c = t.cell(0,0); shade(c, fill); cell_margins(c, 150, 180, 150, 180)
    c.text = ""
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title.upper()); set_font(r, size=9.2, bold=True, color=GOLD)
    p2 = c.add_paragraph(); p2.paragraph_format.space_after = Pt(0); p2.paragraph_format.line_spacing = 1.12
    r2 = p2.add_run(body); set_font(r2, size=10.3, bold=True, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def page_break(): doc.add_page_break()

def add_page_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve"); instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    run._r.extend([fldChar1, instrText, fldChar2])

# Running furniture
header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = header.add_run("ARTDACI  |  Master Production Plan")
set_font(r, size=8.2, bold=True, color=MUTED)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = footer.add_run("Plan directeur · version adaptée au projet · ")
set_font(r, size=8, color=MUTED)
add_page_field(footer)

# Cover
add_p("ARTDACI", bold=True, color=GOLD, align=WD_ALIGN_PARAGRAPH.CENTER, size=13, after=8)
add_p("MASTER PRODUCTION PLAN", bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, size=29, after=6)
add_p("Livre imprimé autonome · Musée augmenté facultatif", color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, size=15, after=26)
add_p("Da Vinci · Vermeer · Van Gogh · Monet", bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, size=12.5, after=46)
callout("Doctrine éditoriale", "Le livre imprimé est l’œuvre de référence. Il doit être complet, utile et durable sans téléphone, sans caméra, sans connexion Internet et sans équipement immersif. L’AR, la VR, la 3D, l’IA, l’audio, la vidéo et les QR codes ajoutent des couches facultatives ; ils ne remplacent jamais le contenu essentiel.")
add_p("Document de pilotage éditorial, scientifique, iconographique, technique et industriel", italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.2, after=8)
add_p("Version du 11 août 2026", color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, size=9.5)
page_break()

heading("1. Résumé exécutif", 1)
add_p("ARTDACI est un projet de livre d’art imprimé consacré à quatre peintres majeurs — Léonard de Vinci, Johannes Vermeer, Vincent van Gogh et Claude Monet — prolongé par un écosystème numérique optionnel. Le présent plan transforme la logique initiale “print + technologies” en une logique éditoriale hiérarchisée : d’abord une référence imprimée indépendante, ensuite des extensions numériques, enfin une plateforme durable capable de les distribuer.")
callout("Décision structurante", "Toute information indispensable à la compréhension des artistes et des œuvres doit figurer dans le livre. Un lecteur qui ne scanne aucun QR code doit recevoir une expérience culturelle complète.", PALE)
heading("1.1 Livrables principaux", 2)
add_bullets([
    "Un livre imprimé de référence, richement illustré, organisé en quatre monographies et un parcours comparatif.",
    "Un corpus éditorial validé : biographies, chronologies, analyses, techniques, contextes, provenances, conservation, réception et bibliographies.",
    "Une iconographie haute définition avec droits, crédits, légendes éditoriales et fichiers maîtres distincts des cibles AR.",
    "Un site multilingue et une galerie VR offrant audio, vidéo, modèles 3D, image tracking, cinéma virtuel, Livre Vivant et guide ChatGPT.",
    "Un système de QR codes et d’URL stables permettant d’atteindre les contenus numériques sans rendre le livre dépendant de ceux-ci.",
])
heading("1.2 Périmètre éditorial", 2)
add_p("Le corpus de production comprend 32 œuvres (huit par peintre). Quatre œuvres signatures reçoivent un traitement prioritaire et exemplaire : Mona Lisa, Girl with a Pearl Earring, Self-Portrait de Van Gogh et Impression, Sunrise. Elles structurent les prototypes imprimés et numériques.")

heading("2. Principes non négociables", 1)
rows = [
    ["Principe", "Exigence de production"],
    ["Print first", "Le texte, les images, les repères historiques et l’analyse nécessaires sont imprimés."],
    ["Autonomie", "Aucun chapitre ne renvoie au numérique pour compléter une lacune essentielle."],
    ["Enrichissement", "Chaque média numérique répond à une question que le papier traite moins bien : mouvement, son, volume, espace ou interaction."],
    ["Durabilité", "Les URL sont stables, les QR remplaçables par redirection et le livre reste valable si un service disparaît."],
    ["Rigueur", "Les faits, datations, attributions, valeurs et interprétations sont sourcés et distingués."],
    ["Accessibilité", "Transcriptions, sous-titres, alternatives textuelles, contraste et navigation sans caméra sont obligatoires."],
    ["Sobriété", "Les modèles, vidéos et textures sont optimisés pour mobile et Meta Quest ; le chargement progressif prévaut."],
]
table(rows, [1.45, 5.3])
page_break()

heading("3. Architecture du livre imprimé", 1)
heading("3.1 Format éditorial recommandé", 2)
add_bullets([
    "Format cible : 23 × 28 cm ou 24 × 30 cm, reliure premium permettant une ouverture confortable.",
    "Pagination cible : 280 à 320 pages, ajustée après chemin de fer et essais d’impression.",
    "Papier intérieur : couché mat ou demi-mat 135–170 g/m² ; épreuves obligatoires pour la colorimétrie, le contraste et les reflets.",
    "Couverture : cartonnée ou souple premium selon coût, avec titre lisible et identité ARTDACI durable, sans promesse technologique dominante.",
])
heading("3.2 Chemin de fer indicatif", 2)
table([
    ["Partie", "Volume cible", "Contenu indispensable"],
    ["Ouverture", "16–20 p.", "Préface, mode d’emploi du livre, chronologie générale, cartes et méthode de lecture."],
    ["Léonard de Vinci", "58–68 p.", "Biographie, Renaissance, science et atelier, huit œuvres, dossier Mona Lisa."],
    ["Johannes Vermeer", "52–62 p.", "Delft, marché de l’art néerlandais, lumière et intérieur, huit œuvres, dossier Pearl Earring."],
    ["Vincent van Gogh", "58–68 p.", "Vie, lettres, Paris/Arles/Saint-Rémy/Auvers, couleur et touche, huit œuvres, dossier Self-Portrait."],
    ["Claude Monet", "58–68 p.", "Impressionnisme, plein air, séries, Giverny, huit œuvres, dossier Impression, Sunrise."],
    ["Regards croisés", "20–28 p.", "Portrait, lumière, espace, couleur, nature, modernité, héritages."],
    ["Références", "14–20 p.", "Glossaire, chronologies, bibliographies, crédits, index, ressources numériques facultatives."],
], [1.4, 1.05, 4.3], font_size=8.8)
heading("3.3 Gabarit de chapitre par peintre", 2)
add_numbers([
    "Portrait intellectuel et biographique : formation, lieux, réseaux, périodes et tournants.",
    "Contexte historique, social, culturel et technique.",
    "Méthodes de travail : supports, préparation, dessin, pigments, touche, lumière et composition.",
    "Chronologie illustrée et cartographie des lieux de vie et de création.",
    "Huit notices d’œuvres, dont une œuvre signature traitée sur plusieurs doubles pages.",
    "Réception, marché, conservation, influence et héritage culturel.",
    "Bibliographie sélective et sources institutionnelles.",
])
heading("3.4 Gabarit obligatoire de notice d’œuvre", 2)
add_bullets([
    "Reproduction principale de qualité, détails pertinents et données d’identification complètes.",
    "Date, support, technique, dimensions, collection, provenance et état des connaissances sur l’attribution.",
    "Contexte de commande ou de création ; place dans la carrière du peintre.",
    "Analyse guidée de la composition, de la lumière, de la couleur, de l’espace, des gestes et des motifs.",
    "Techniques matérielles et résultats de recherche scientifique, avec niveau de certitude explicite.",
    "Histoire de la réception, conservation, déplacements, vols ou restaurations lorsque pertinents.",
    "Valeur culturelle et, si une valeur financière est évoquée, date, nature de l’estimation et prudence méthodologique.",
    "Encadré facultatif “Prolonger l’expérience” regroupant QR, AR, audio, vidéo ou 3D sans interrompre la lecture.",
])

heading("4. Les quatre dossiers signatures", 1)
signature_rows = [
    ["Œuvre", "Dossier imprimé autonome", "Extension facultative"],
    ["Mona Lisa", "Portrait, commande probable, sfumato, regard et sourire, paysage, mains, matérialité, vol de 1911, Louvre, conservation et célébrité.", "Image tracking MindAR, modèles GLB, animation MP4, audio multilingue, hotspots, studio de Léonard et galerie VR."],
    ["Girl with a Pearl Earring", "Tronie, Delft vers 1665, pose, lumière, bleu ultramarin, matière de la perle, conservation et réception moderne.", "Cible MindAR, modèles portrait/assis, animation, audio, 3D et salle Vermeer."],
    ["Self-Portrait de Van Gogh", "Contexte biographique précis, autoportrait comme laboratoire, touche, palette, regard, état matériel et série des autoportraits.", "Cible MindAR, modèle 3D, animation, audio et salle Van Gogh ; liens vers The Bedroom."],
    ["Impression, Sunrise", "Le Havre, 1872, exposition de 1874, naissance du terme impressionnisme, complémentarité orange/bleu, atmosphère et modernité industrielle.", "Animation préparée, audio et salle Monet ; cible MindAR à produire et valider avant activation."],
]
table(signature_rows, [1.45, 3.05, 2.25], font_size=8.4)

heading("5. Corpus de référence : 32 œuvres", 1)
works = [
    ("LD01","Mona Lisa","Signature"),("LD02","Lady with an Ermine","Haute"),("LD03","Ginevra de’ Benci","Haute"),("LD04","La Belle Ferronnière","Standard"),("LD05","The Annunciation","Haute"),("LD06","Virgin of the Rocks","Haute"),("LD07","The Last Supper","Haute"),("LD08","Saint John the Baptist","Standard"),
    ("VE01","Girl with a Pearl Earring","Signature"),("VE02","The Milkmaid","Haute"),("VE03","View of Delft","Haute"),("VE04","Woman Holding a Balance","Standard"),("VE05","The Art of Painting","Haute"),("VE06","The Astronomer","Standard"),("VE07","The Geographer","Standard"),("VE08","Woman in Blue Reading a Letter","Standard"),
    ("VG01","The Starry Night","Haute"),("VG02","Sunflowers","Haute"),("VG03","The Bedroom","Haute"),("VG04","Café Terrace at Night","Standard"),("VG05","The Night Café","Standard"),("VG06","Self-Portrait","Signature"),("VG07","Irises","Haute"),("VG08","Almond Blossom","Standard"),
    ("MO01","Impression, Sunrise","Signature"),("MO02","Woman with a Parasol","Standard"),("MO03","Bridge at Argenteuil","Haute"),("MO04","Poppies","Standard"),("MO05","The Japanese Bridge","Haute"),("MO06","Water Lilies","Haute"),("MO07","Rouen Cathedral Series","Haute"),("MO08","Houses of Parliament Series","Haute"),
]
rows = [["Code","Œuvre","Traitement imprimé","Extension numérique recommandée"]]
for code, title, priority in works:
    ext = "Notice complète + iconographie + sources"
    if priority == "Signature": ext = "Dossier signature + AR/animation/audio/3D selon disponibilité"
    elif priority == "Haute": ext = "Notice approfondie + média ciblé si pertinent"
    rows.append([code,title,priority,ext])
table(rows, [0.58, 2.22, 1.08, 2.87], font_size=8.15)

heading("6. Écosystème numérique facultatif", 1)
add_p("Le numérique est conçu comme un musée augmenté distribué autour du livre, non comme une condition d’accès au savoir. Chaque extension doit annoncer clairement sa valeur ajoutée et proposer un retour simple à la lecture.")
table([
    ["Couche", "Fonction", "Technologie observée dans ARTDACI", "Alternative sans technologie"],
    ["QR / Web", "Accès stable à une ressource", "Pages HTML, paramètres painting/lang, QR vers URL", "URL courte imprimée et contenu essentiel dans le livre"],
    ["Image tracking", "Reconnaître une reproduction imprimée", "MindAR, cible .mind, getUserMedia, HTTPS", "Bouton/QR d’accès manuel et notice imprimée complète"],
    ["3D", "Montrer volume, objet ou espace", "Three.js, GLTFLoader, GLB/glTF, Draco", "Vues fixes, schémas, plans et détails imprimés"],
    ["AR spatiale", "Placer un modèle dans l’espace", "model-viewer et capacités appareils selon support", "Photomontages et échelles dans le livre"],
    ["VR", "Entrer dans les salles et mondes", "WebXR, Meta Quest Browser, contrôleurs, hand tracking", "Parcours visuel et descriptions imprimées"],
    ["Audio / vidéo", "Narration, mouvement et séquence", "MP3/M4A/MP4, textures vidéo, synthèse vocale de secours", "Transcription, photogrammes et texte"],
    ["IA", "Guide conversationnel contextuel", "Lien ChatGPT prérempli, contexte de salle et langue", "Index, glossaire, FAQ et notices imprimées"],
], [1.0, 1.3, 2.5, 1.95], font_size=7.9)
heading("6.1 Expériences actuellement représentées dans le dépôt", 2)
add_bullets([
    "Galerie VR connectée avec quatre salles de peinture, salles de modèles 3D, groupes réimaginés, musée du Louvre, cinéma et téléportation.",
    "Compatibilité WebXR, interaction par contrôleurs et pincement des mains, avec rendu adapté aux navigateurs Meta Quest.",
    "Image tracking pour Mona Lisa, Van Gogh, The Bedroom et Girl with a Pearl Earring ; animations dédiées pour les quatre œuvres signatures.",
    "Livre Vivant 3D, audio multilingue, vidéos, variantes de modèles et interfaces français/anglais/arabe.",
    "Guide ChatGPT accessible depuis l’interface et certains panneaux de navigation, avec question et contexte préremplis.",
])
heading("6.2 État technique des quatre œuvres signatures", 2)
table([
    ["Œuvre", "JSON", "Cible .mind", "Vidéo AR", "Statut"],
    ["Mona Lisa", "Oui", "Oui", "Oui", "Intégrée ; QA mobile/Quest requise"],
    ["Self-Portrait", "Oui", "Oui", "Oui", "Intégrée ; QA caméra et lisibilité"],
    ["Girl with a Pearl Earring", "Oui", "Oui", "Oui", "Intégrée ; QA tracking fond sombre"],
    ["Impression, Sunrise", "Oui", "À produire", "Oui", "Animation prête ; tracking non activable avant cible"],
], [1.45, 0.75, 1.05, 1.0, 2.5], font_size=8.5)

heading("7. Architecture technique de référence", 1)
heading("7.1 Composants", 2)
add_bullets([
    "Front-end statique : HTML5, CSS, JavaScript en modules ES, données éditoriales en JSON.",
    "Moteur 3D : Three.js ; chargement GLB/glTF ; décodage Draco pour les modèles compressés.",
    "WebAR : MindAR Image Tracking avec cibles compilées .mind et autorisation caméra via getUserMedia.",
    "VR : WebXR avec local-floor/bounded-floor, contrôleurs et fonctionnalité optionnelle hand-tracking.",
    "Multimédia : images WebP/PNG/JPEG, audio MP3/M4A, vidéo MP4 et textures vidéo dans la scène 3D.",
    "Compatibilité : chargement progressif, baisse de résolution de textures sur mobile/Quest, modèles allégés et contenu différé.",
    "Internationalisation : interfaces et contenus en français, anglais et arabe lorsque disponibles.",
])
heading("7.2 Architecture de données", 2)
add_p("Une œuvre est pilotée par un manifeste JSON qui sépare identité, textes, sources, impression, AR, modèles, audio, vidéo, hotspots et objectifs éducatifs. Cette séparation permet de corriger un texte ou remplacer un média sans reconstruire toute l’application.")
table([
    ["Objet", "Contenu minimal", "Règle"],
    ["painting.json", "slug, titre, artiste, date, technique, collection, textes, sources", "Source éditoriale versionnée"],
    ["print", "ordre, image cible, QR, crédits", "Indépendant des médias lourds"],
    ["ar", "cible .mind, modèle principal, variantes, échelle, rotation", "Pas de chemin publié sans fichier présent"],
    ["media", "image, audio, vidéo, modèles", "Formats Web optimisés et crédits associés"],
    ["education", "objectifs, vocabulaire, quiz, questions", "Cohérent avec le texte imprimé"],
], [1.15, 3.35, 2.25], font_size=8.6)
heading("7.3 Performance et robustesse", 2)
add_bullets([
    "Budget initial minimal : HTML, styles, scripts et interface avant les gros médias.",
    "Chargement à la demande des salles, modèles et pistes audio ; pas de préchargement intégral sur mobile ou Quest.",
    "Cibles AR et visuels de suivi distincts des images maîtres d’impression.",
    "Textures limitées selon l’appareil ; mipmaps désactivées lorsque pertinent ; modèles low-power pour les scènes lourdes.",
    "Mesure obligatoire : temps d’ouverture, mémoire, FPS, délai de reconnaissance, erreurs caméra et redémarrages de page.",
])

heading("8. Gouvernance éditoriale et scientifique", 1)
heading("8.1 Chaîne de validation", 2)
add_numbers([
    "Constituer le dossier source institutionnel et bibliographique de chaque œuvre.",
    "Rédiger une fiche longue factuelle, puis la notice imprimée et les scripts numériques dérivés.",
    "Identifier les affirmations certaines, discutées, hypothétiques ou dépassées.",
    "Faire relire par un spécialiste ou un comité éditorial compétent.",
    "Valider les crédits, droits, dates d’accès et licences des reproductions.",
    "Geler une version imprimée ; poursuivre les mises à jour numériques sans altérer la cohérence du livre.",
])
heading("8.2 Politique sur la valeur des œuvres", 2)
callout("Prudence", "Ne pas présenter une estimation spéculative comme un prix. Pour une œuvre inaliénable de musée, expliquer l’absence de marché direct ; distinguer valeur culturelle, valeur assurantielle éventuelle, records comparables et date de l’information.", PALE)
heading("8.3 Sources prioritaires", 2)
add_p("Privilégier les musées détenteurs, catalogues raisonnés, publications scientifiques, dossiers de conservation et bases institutionnelles. Les ressources généralistes servent à l’orientation, jamais seules à une affirmation sensible.")

heading("9. Droits, crédits et gestion des actifs", 1)
table([
    ["Statut", "Définition", "Action"],
    ["R0", "Non vérifié", "Interdiction de publication"],
    ["R1", "Œuvre dans le domaine public", "Vérifier séparément le fichier de reproduction"],
    ["R2", "Reproduction Open Access", "Conserver licence et crédit exact"],
    ["R3", "Crédit obligatoire", "Insérer crédit dans livre et métadonnées"],
    ["R4", "Autorisation requise", "Obtenir accord écrit"],
    ["R5", "Usage commercial à négocier", "Bloquer production finale jusqu’au contrat"],
], [0.65, 3.0, 3.1], font_size=8.6)
add_bullets([
    "Séparer les droits de l’œuvre, de la reproduction photographique, du modèle 3D, de l’enregistrement sonore et de la vidéo.",
    "Conserver source originale, fichier de travail, master print, version Web et preuve de licence.",
    "Ne jamais écraser un original ; utiliser code œuvre, type d’actif, langue et version dans chaque nom de fichier.",
    "Associer à chaque actif : créateur, licence, date, source, restrictions, crédit et date d’expiration éventuelle.",
])

heading("10. Production iconographique et impression", 1)
heading("10.1 Image maître", 2)
add_bullets([
    "Résolution effective de 300 dpi à la taille d’impression et profil colorimétrique défini avec l’imprimeur.",
    "Contrôle du cadrage, des noirs, de la saturation, des dominantes, du piqué et des détails agrandis.",
    "Épreuve contractuelle pour les quatre œuvres signatures et les images particulièrement sombres ou saturées.",
    "Fichier AR_TARGET séparé, optimisé pour navigateur et évalué pour sa richesse en points caractéristiques.",
])
heading("10.2 QR codes et signalétique", 2)
add_bullets([
    "URL courte et durable sous chaque QR ; aucune dépendance à un chemin de fichier interne.",
    "Zone calme suffisante, contraste élevé, taille testée sur épreuve réelle et correction d’erreur adaptée.",
    "Libellé explicite : “Écouter”, “Voir l’animation”, “Explorer en 3D” ou “Entrer dans la galerie VR”.",
    "Les icônes ne doivent ni masquer le texte ni concurrencer l’œuvre ; elles se regroupent dans une marge cohérente.",
])

heading("11. Accessibilité, vie privée et éthique", 1)
add_bullets([
    "Le livre imprimé emploie une typographie lisible, un contraste suffisant, une structure répétable et des légendes complètes.",
    "Chaque audio dispose d’une transcription ; chaque vidéo, de sous-titres ; chaque image numérique, d’un texte alternatif pertinent.",
    "Toutes les fonctions essentielles restent accessibles sans AR, sans VR, sans gestes et sans activation de la caméra.",
    "Le flux caméra sert localement à la reconnaissance d’image lorsque MindAR fonctionne dans le navigateur ; cette finalité doit être expliquée.",
    "Le guide IA doit signaler sa nature, éviter de présenter une interprétation comme un fait et renvoyer aux sources éditoriales.",
    "Les statistiques doivent être minimales, anonymisées et proportionnées ; aucun suivi invasif n’est nécessaire au fonctionnement du livre.",
])

heading("12. Plan de production", 1)
heading("12.1 Lots", 2)
table([
    ["Lot", "Objectif", "Sortie validée"],
    ["0 — Gouvernance", "Charte éditoriale, droits, nomenclature, modèles de notice", "Kit de production approuvé"],
    ["1 — Quatre signatures", "Dossiers imprimés complets et extensions existantes", "Prototype représentatif"],
    ["2 — Monographies", "Biographies, chronologies et 28 notices restantes", "Manuscrit complet"],
    ["3 — Iconographie", "Masters, détails, cartes, schémas, crédits", "Dossier image verrouillé"],
    ["4 — Design éditorial", "Chemin de fer, maquette, index, références", "Bon à tirer"],
    ["5 — Numérique", "QR, Web, AR, audio, vidéo, 3D, VR", "Extensions publiées et documentées"],
    ["6 — QA / lancement", "Épreuves, appareils, performance, accessibilité", "Publication print et Web"],
], [1.2, 3.05, 2.5], font_size=8.5)
heading("12.2 Ordre par œuvre", 2)
add_numbers([
    "Droits et sources.", "Texte long validé.", "Notice imprimée et iconographie.", "Maquette et épreuve.",
    "URL/QR stable.", "Audio et vidéo dérivés du texte validé.", "AR et 3D si valeur ajoutée démontrée.",
    "QA multi-appareil et archivage de version.",
])
heading("12.3 Statuts recommandés", 2)
table([
    ["Statut", "Sens"],
    ["BACKLOG", "Prévu, non commencé"], ["IN PROGRESS", "Production active"],
    ["REVIEW", "Relecture scientifique, juridique ou technique"], ["BLOCKED", "Dépendance explicite non résolue"],
    ["VALIDATED", "Contenu ou actif approuvé"], ["PUBLISHED", "Version publique et archivée"],
], [1.55, 5.2], font_size=8.8)

heading("13. Contrôle qualité", 1)
heading("13.1 Livre imprimé", 2)
add_bullets([
    "Exactitude des textes, cohérence des dates, noms, titres, dimensions et collections.",
    "Qualité des reproductions, détails non pixellisés, crédits présents, profils colorimétriques maîtrisés.",
    "Absence de veuves/orphelines gênantes, légendes séparées, tableaux lisibles, index vérifié et renvois corrects.",
    "Lecture complète sans QR : aucune phrase essentielle du type “scannez pour comprendre”.",
])
heading("13.2 Numérique", 2)
table([
    ["Plateforme", "Tests minimum"],
    ["Android / Chrome", "Caméra, QR, image tracking, audio/vidéo, mémoire, orientation"],
    ["iPhone / Safari", "Permission caméra, lecture inline, gestes, retour arrière, Quick Look si utilisé"],
    ["Meta Quest Browser", "WebXR, téléportation, contrôleurs, mains, FPS, chargement progressif"],
    ["Ordinateur", "Navigation clavier/souris, contenu de secours, redimensionnement"],
    ["Réseau lent", "Affichage initial, chargement différé, erreurs compréhensibles, reprise"],
], [1.5, 5.25], font_size=8.7)
heading("13.3 Seuils pilotes", 2)
add_bullets([
    "QR : 100 % des liens valides sur l’épreuve finale.",
    "Tracking : reconnaissance répétable sous plusieurs lumières et distances, avec solution manuelle de secours.",
    "Stabilité : aucun redémarrage de page durant un parcours standard sur téléphones de référence et Quest.",
    "Temps vers premier contenu utile : cible inférieure à 5 secondes sur connexion moyenne, hors média lourd explicite.",
    "Accessibilité : transcriptions, sous-titres, textes alternatifs, contraste et navigation de secours validés.",
])

heading("14. Risques et réponses", 1)
table([
    ["Risque", "Impact", "Réponse"],
    ["Le numérique domine le livre", "Obsolescence et lecture fragmentée", "Geler d’abord le manuscrit autonome et limiter les appels numériques"],
    ["Droits incomplets", "Retard ou retrait d’images", "Registre des droits, alternatives Open Access et gate juridique"],
    ["Modèles trop lourds", "Lag, crash, redémarrage mobile", "Draco, LOD, textures réduites, lazy loading, budget par salle"],
    ["Tracking instable", "Échec utilisateur", "Cible testée sur papier, éclairage varié, QR et lancement manuel"],
    ["Lien imprimé obsolète", "Livre amputé de ses bonus", "URL courte contrôlée, redirection et page de secours"],
    ["Contenu IA inexact", "Perte de confiance", "Prompt cadré, sources, avertissement et validation éditoriale"],
    ["Incohérences multilingues", "Expérience inégale", "Glossaire, mémoire de traduction et QA par langue"],
], [1.55, 2.15, 3.05], font_size=8.25)

heading("15. Matrice de validation finale", 1)
checks = [
    ("Éditorial", "Les quatre monographies et 32 notices sont complètes, sourcées et relues."),
    ("Autonomie", "Le livre reste pleinement utile sans accès numérique."),
    ("Iconographie", "Masters, colorimétrie, détails, légendes et crédits sont validés."),
    ("Droits", "Chaque actif possède un statut et une preuve exploitable."),
    ("Print", "BAT, papier, reliure, contraste, QR et index sont approuvés."),
    ("AR", "Cibles présentes, testées sur papier, avec fallback manuel."),
    ("3D / VR", "Échelles naturelles, orientation, performance, navigation et accessibilité sont testées."),
    ("Audio / vidéo", "Langues, niveaux, sous-titres, transcriptions et crédits sont complets."),
    ("IA", "Guide identifié, contextualisé et limité aux usages d’accompagnement."),
    ("Archivage", "Sources, masters, livrables, versions et changelog sont conservés."),
]
table([["Domaine","Critère de sortie"]] + checks, [1.25, 5.5], font_size=8.6)

page_break()
heading("16. Prochaines décisions", 1)
add_numbers([
    "Valider le sommaire éditorial définitif et la pagination cible du livre.",
    "Désigner les responsables de validation scientifique, iconographique, juridique et technique.",
    "Terminer les quatre dossiers signatures imprimés avant d’élargir les expériences numériques.",
    "Produire et tester la cible MindAR d’Impression, Sunrise, sans publier de chemin incomplet.",
    "Établir un budget de poids par page Web, modèle GLB, texture, vidéo et salle VR.",
    "Réaliser un prototype papier de 32–48 pages et le tester comme livre, indépendamment de l’application.",
    "Après validation du prototype, lancer les 28 notices restantes par lots éditoriaux.",
])
callout("Critère de réussite", "ARTDACI est réussi si le livre peut devenir une référence durable sur les quatre peintres, tandis que ses extensions numériques donnent envie d’aller plus loin sans jamais être nécessaires pour comprendre.")

# Core properties
doc.core_properties.title = "ARTDACI Master Production Plan amélioré"
doc.core_properties.subject = "Plan directeur du livre imprimé autonome et de ses extensions numériques facultatives"
doc.core_properties.author = "ARTDACI"
doc.core_properties.keywords = "ARTDACI, livre d'art, Da Vinci, Vermeer, Van Gogh, Monet, AR, VR, 3D, QR, IA"

doc.save(OUT)
print(OUT)
